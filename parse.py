import sys
import os
import traceback
import json
import jsonlines
from docx import Document
from card import Card, TAG_NAME
import argparse
from datasets import load_dataset, ClassLabel, Sequence

from utils import format_prompt_for_openai_completion

CITE_NAME = "13 pt Bold"

def parse_cites(filename):
  document = Document(filename)
  cites = []
  print("Parsing " + filename)
  print("Found " + str(len(document.paragraphs)) + " paragraphs")

  for paragraph in document.paragraphs:
    cite = paragraph.text
    for r in paragraph.runs:
      if CITE_NAME in r.style.name or (r.style.font.bold or r.font.bold):
        cite = cite.replace(r.text, "**" + r.text + "**")
    cites.append(cite)
  
  return cites

def parse_cards(filename):
  document = Document(filename)
  cards = []
  current_card = []
  print("Parsing " + filename)
  for paragraph in document.paragraphs:
    if paragraph.style.name == TAG_NAME:
      try:
        cards.append(Card(current_card))
      except Exception as e:
        continue
      finally:
        current_card = [paragraph]
    else:
      current_card.append(paragraph)
  
  return cards

def push_to_hf(output_file, hub_name):
  dataset = load_dataset("json", data_files=output_file)
  features = dataset['train'].features.copy()

  features["highlight_labels"] = Sequence(feature=ClassLabel(names=["no", "yes"]), id=None)
  features["underline_labels"] = Sequence(feature=ClassLabel(names=["no", "yes"]), id=None)
  features["emphasis_labels"] = Sequence(feature=ClassLabel(names=["no", "yes"]), id=None)

  dataset = dataset['train'].map(lambda x : x, features=features, batched=True)
  dataset.push_to_hub(hub_name)

if __name__ == "__main__":
  # Parse command line arguments
  parser = argparse.ArgumentParser()
  parser.add_argument("file", type=str, help="path to a docx file or directory of files")
  parser.add_argument("-o", "--output", type=str, help="path to output file (default: output.json)", default="output.jsonl")
  parser.add_argument("-jsonl", "--jsonl", help="force output in jsonl format (default: False)", action='store_true')
  parser.add_argument("-s", "--skip_parse", help="skip parsing and use existing output file", action='store_true')
  parser.add_argument("-hub", "--hub", help="provide to format as a huggingface dataset and push output (must be logged in)", type=str, default=None)
  parser.add_argument("--field", help="which aspect of the card to predict (highlights, underlines, emphasis)", type=str, default="highlights")
  parser.add_argument("--input_field", help="which aspect of the card to use as input (text, underlines)", type=str, default="text")
  parser.add_argument("--hub_format", help="force save in hf format", action='store_true')

  args = parser.parse_args()

  # Upload existing output file to huggingface hub
  if args.skip_parse:
    if not os.path.isfile(args.output):
      print("Output file not found")
      sys.exit(1)
    if not args.hub:
      print("No hub name provided")
      sys.exit(1)
    push_to_hf(args.output, args.hub)
    sys.exit(0)

  # Get list of files to parse
  files_to_parse = []
  file_or_dir = args.file
  if os.path.isfile(file_or_dir):
    files_to_parse.append(file_or_dir)
  elif os.path.isdir(file_or_dir):
    for file in os.listdir(file_or_dir):
      if file.endswith(".docx"):
        files_to_parse.append(os.path.join(file_or_dir, file))
  else:
    print("File not found")
    sys.exit(1)
  
  # Parse each file into a list of cards
  cards = []
  for docx_name in files_to_parse:
    try:
      parsed_cards = parse_cards(docx_name)
      cards.extend(parsed_cards)
    except Exception as e:
      print("Error parsing " + docx_name)
      continue

  print("Found " + str(len(cards)) + " cards")

  # Strip punctuation and empty strings from each card's highlighted text
  for card in cards:
    punctuation_list = [",", ".", "!", "?", ":", ";", "(", ")", "[", "]", "{", "}", "\"", "\'", "“", "”", "‘", "’"]

    # Strip punctuation from each word in list of highlighted words
    card.highlighted_text = [word.strip("".join(punctuation_list)) for word in card.highlighted_text]
    card.underlined_text = [word.strip("".join(punctuation_list)) for word in card.underlined_text]
    card.emphasized_text = [word.strip("".join(punctuation_list)) for word in card.emphasized_text]

    # Remove empty strings
    card.highlighted_text = list(filter(None, card.highlighted_text))
    card.underlined_text = list(filter(None, card.underlined_text))
    card.emphasized_text = list(filter(None, card.emphasized_text))

    # Assert that run_text length == highlight/underline/emphasis length
    try:
      assert len(card.run_text) == len(card.highlight_labels) 
      assert len(card.run_text) == len(card.underline_labels) 
      assert len(card.run_text) == len(card.emphasis_labels)
    except AssertionError:
      print("Error parsing " + card.tag + ": run_text length does not match highlight/underline/emphasis length")

    # Remove empty strings from run_text (and the associated labels)
    # Keep track of indexes to remove from labels
    indexes_to_remove = []
    for i, word in enumerate(card.run_text):
      if word == "":
        indexes_to_remove.append(i)
    for i in sorted(indexes_to_remove, reverse=True):
      del card.run_text[i]
      del card.highlight_labels[i]
      del card.underline_labels[i]
      del card.emphasis_labels[i]

  # Strip \u2018 and \u2019 and \u2014 from each card's card_text, tag, and underlined_text
  for card in cards:
    card.card_text = card.card_text.replace("\u2018", "'").replace("\u2019", "'").replace("\u2014", "-")
    card.tag = card.tag.replace("\u2018", "'").replace("\u2019", "'").replace("\u2014", "-")
    card.underlined_text = [word.replace("\u2018", "'").replace("\u2019", "'").replace("\u2014", "-") for word in card.underlined_text]
    card.highlighted_text = [word.replace("\u2018", "'").replace("\u2019", "'").replace("\u2014", "-") for word in card.highlighted_text]
    card.emphasized_text = [word.replace("\u2018", "'").replace("\u2019", "'").replace("\u2014", "-") for word in card.emphasized_text]
    card.run_text = [word.replace("\u2018", "'").replace("\u2019", "'").replace("\u2014", "-") for word in card.run_text]
    
  # Write cards to JSON file
  json_dict = [{
    "tag": card.tag, 
    "text": card.card_text, 
    "highlights": card.highlighted_text, 
    "underlines": card.underlined_text,
    "emphasis": card.emphasized_text,
    "cite": card.cite,
    "cite_emphasis": card.cite_emphasis,
    "run_text": card.run_text,
    "highlight_labels": card.highlight_labels,
    "underline_labels": card.underline_labels,
    "emphasis_labels": card.emphasis_labels,
  } for card in cards]

  output_file = args.output

  if args.hub or args.hub_format:
    if output_file.endswith(".json") and not args.jsonl:
      with open(output_file, "w") as outfile:
        print("Writing to " + output_file)
        json.dump(json_dict, outfile, indent=4)
    elif output_file.endswith(".jsonl") or args.jsonl:
      with jsonlines.open(output_file, mode="w") as writer:
        print("Writing to " + output_file)
        writer.write_all([card for card in json_dict])
    else:
      print("Invalid output file type")
      sys.exit(1)

    if args.hub:
      dataset = load_dataset("json", data_files=output_file)
      push_to_hf(output_file, args.hub)
  else:
    with jsonlines.open(output_file, mode="w") as writer:
      print("Writing to " + output_file)

      if args.input_field == "text":
        for card in json_dict:
          prompts, _ = format_prompt_for_openai_completion(card["tag"], card["text"])
          writer.write_all([{
            "prompt": prompt,
            "completion": " " + json.dumps(card[args.field]) + " END"
          } for prompt in prompts])
      else:
        for card in json_dict:
          prompts, _ = format_prompt_for_openai_completion(card["tag"], json.dumps(card[args.input_field]))
          writer.write_all([{
            "prompt": prompt,
            "completion": " " + json.dumps(card[args.field]) + " END"
          } for prompt in prompts])

    # with jsonlines.open("emphasis.jsonl", mode="w") as writer:
    #   for card in json_dict:
    #     prompts = format_prompt_for_openai_completion(card["tag"], json.dumps(card[args.input_field]))
    #     writer.write_all([{
    #       "prompt": prompt,
    #       "completion": " " + json.dumps(card["emphasis"]) + " END"
    #     } for prompt in prompts])